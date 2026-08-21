"""
DOCX export service.

Builds a downloadable Word document from a generated StoryData payload.
"""

from io import BytesIO
import base64
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.schemas.response import StoryData


def _safe_filename(title: str) -> str:
    """Return a conservative filename derived from a story title."""
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title.strip().lower()).strip("-")
    return slug[:80] or "generated-story"


def _configure_document_styles(document: Document) -> None:
    """Apply a clean, readable style baseline for generated story documents."""
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11.5, RGBColor(68, 68, 68)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def _add_metadata_block(document: Document, story: StoryData) -> None:
    document.add_heading(story.title, level=1)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle.add_run(story.subtitle)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    source = document.add_paragraph()
    source_run = source.add_run(f"Source: {story.source_context}")
    source_run.bold = True

    count = document.add_paragraph()
    count.add_run(f"Word count: {story.word_count}")


def _add_cover_image(document: Document, story: StoryData) -> None:
    """Embed the generated cover image, if present. Silently does nothing
    if no image was generated — image generation is best-effort, so its
    absence is a normal, non-error state."""
    if not story.cover_image_base64:
        return

    try:
        image_bytes = base64.b64decode(story.cover_image_base64)
    except Exception:
        # Malformed base64 shouldn't break the whole document — just skip
        # the image rather than failing the export.
        return

    document.add_picture(BytesIO(image_bytes), width=Inches(6.2))
    document.add_paragraph()


def _add_summary_table(document: Document, story: StoryData) -> None:
    if not story.summary_table:
        return

    document.add_heading("Key Facts", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, label in enumerate(["Pillar", "Metric", "Purpose"]):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(label)
        run.bold = True

    for item in story.summary_table:
        row_cells = table.add_row().cells
        row_cells[0].text = item.pillar
        row_cells[1].text = item.metric
        row_cells[2].text = item.purpose

    document.add_paragraph()


def _add_markdown_story(document: Document, markdown: str) -> None:
    document.add_heading("Story", level=2)

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("### "):
            document.add_heading(line.removeprefix("### ").strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line.removeprefix("## ").strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line.removeprefix("# ").strip(), level=2)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)


def build_story_docx(story: StoryData) -> tuple[BytesIO, str]:
    """Build a Word document for the provided story and return it as bytes."""
    document = Document()
    _configure_document_styles(document)
    _add_metadata_block(document, story)
    _add_cover_image(document, story)
    _add_summary_table(document, story)
    _add_markdown_story(document, story.story_markdown)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output, f"{_safe_filename(story.title)}.docx"
